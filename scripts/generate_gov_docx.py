#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
党政机关公文生成器 - 基于 GB/T 9704-2012 国家标准
生成符合规范的党政机关公文 docx 文件

支持 4 种公文格式:
  - general: 通用格式 (下行文/平行文，机关标志居中)
  - letter:  信函格式 (上距版心30mm，下距20mm各印红色双线，无发文字号)
  - order:   命令(令)格式 (机关标志全称+「命令」/「令」，独占编排令号)
  - minutes: 纪要格式 ("×××××纪要"标志 + 出席人员名单)
"""

import json
import sys
import argparse
import os
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from copy import deepcopy


# ============================================================
# 字号对照表 (中文字号 -> 磅值)
# ============================================================
FONT_SIZE_MAP = {
    '初号': 42, '小初': 36, '一号': 26, '小一': 24,
    '二号': 22, '小二': 18, '三号': 16, '小三': 15,
    '四号': 14, '小四': 12, '五号': 10.5, '小五': 9,
}

# 结构层次对应的字体
LEVEL_FONT_MAP = {
    1: '黑体',    # 一、二、三、
    2: '楷体',    # （一）（二）（三）
    3: '仿宋',    # 1. 2. 3.
    4: '仿宋',    # （1）（2）（3）
}

# 各种公文格式
FORMAT_TYPES = ('general', 'letter', 'order', 'minutes')

# 六角括号 (GB/T 9704 §3.1.5 要求用 "〔〕")
BRACKET_OPEN = '〔'   # 〔
BRACKET_CLOSE = '〕'  # 〕


def normalize_brackets(year):
    """将 ASCII 方括号 / 全角方括号 / 已包六角括号 归一为只含数字。

    返回纯数字 (无任何括号)；调用方负责重新包成 〔...〕。
    """
    if year is None:
        return ''
    s = str(year).strip()
    # 移除任意括号包围 (ASCII / 全角 / 六角)
    for lo, hi in [('[', ']'), ('［', '］'), (BRACKET_OPEN, BRACKET_CLOSE)]:
        s = s.strip(lo + hi)
    return s


def format_doc_number(prefix, year, sequence):
    """生成 GB/T 9704 §3.1.5 标准的发文字号：「机关代字〔年份〕序号号」"""
    y = normalize_brackets(year)
    return f'{prefix or ""}{BRACKET_OPEN}{y}{BRACKET_CLOSE}{sequence or ""}号'


def set_font(run, font_name, font_size_pt, bold=False, color=None, east_asia_font=None):
    """设置run的字体属性"""
    run.font.name = font_name
    run.font.size = Pt(font_size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    # 设置中文字体
    ea = east_asia_font or font_name
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{ea}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), ea)


def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=None, line_rule=None):
    """设置段落间距"""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)


def add_empty_lines(doc, count=1, font_size_pt=16):
    """添加空行"""
    for _ in range(count):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=font_size_pt * 1.75)
        run = p.add_run('')
        set_font(run, '仿宋', font_size_pt)


def setup_page(doc, format_type='general'):
    """设置A4页面，符合GB/T 9704-2012要求

    - general/letter/minutes: 天头 37mm，订口 28mm
    - order (命令(令)):      机关标志上距版心 20mm，天头约 22mm
    """
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.orientation = WD_ORIENT.PORTRAIT
    if format_type == 'order':
        # 命令(令): 机关标志上边缘至版心上边缘 20mm
        section.top_margin = Mm(22)
        section.bottom_margin = Mm(35)
        section.left_margin = Mm(28)
        section.right_margin = Mm(26)
    else:
        # general / letter / minutes: 默认天头 37mm
        section.top_margin = Mm(37)
        section.bottom_margin = Mm(35)
        section.left_margin = Mm(28)
        section.right_margin = Mm(26)
    return section


def _set_paragraph_border(p, color='000000', sz=8, position='bottom'):
    """给段落添加边框（用于版头/版记分隔线）。"""
    pPr = p._element.get_or_add_pPr()
    # 移除已有边框
    existing = pPr.find(qn('w:pBdr'))
    if existing is not None:
        pPr.remove(existing)
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:{position} w:val="single" w:sz="{sz}" '
        f'    w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_red_line(doc, width_mm=156, thickness_pt=1.5):
    """添加红色分隔线（版头分隔线）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0)
    _set_paragraph_border(p, color='FF0000', sz=int(thickness_pt * 8))
    run = p.add_run('')
    run.font.size = Pt(2)
    return p


def add_red_double_line_top(doc, thick_sz=18, thin_sz=10):
    """版头红色双线 (上粗下细) — 用于信函格式 (距上页边 30mm 处)

    GB/T 9704 §4.1: 「发文机关标志下4mm处印红色双线(上粗下细)」
    用一个段落 + 上下两条 w:bottom+top 边框实现。
    """
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{thick_sz}" w:space="1" w:color="FF0000"/>'
        f'  <w:bottom w:val="single" w:sz="{thin_sz}" w:space="1" w:color="FF0000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run('')
    run.font.size = Pt(2)
    return p


def add_red_double_line_bottom(doc, thick_sz=18, thin_sz=10):
    """版尾红色双线 (上细下粗) — 用于信函格式 (距下页边 20mm 处)

    GB/T 9704 §4.1: 「距下页边20mm处印红色双线(上细下粗)」
    """
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{thin_sz}" w:space="1" w:color="FF0000"/>'
        f'  <w:bottom w:val="single" w:sz="{thick_sz}" w:space="1" w:color="FF0000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    run = p.add_run('')
    run.font.size = Pt(2)
    return p


def add_thick_line(doc):
    """添加粗分隔线（版记首条/末条，0.35mm）"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    _set_paragraph_border(p, color='000000', sz=10)
    run = p.add_run('')
    run.font.size = Pt(2)
    return p


def add_thin_line(doc):
    """添加细分隔线（版记中间，0.25mm）"""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    _set_paragraph_border(p, color='000000', sz=6)
    run = p.add_run('')
    run.font.size = Pt(2)
    return p


# ============================================================
# 版头构建 (4 种格式分别 dispatch)
# ============================================================

def _add_copy_number(doc, header_data):
    """份号 (顶格版心左上角第一行) — 4 种格式共有"""
    copy_number = header_data.get('copy_number')
    if not copy_number:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(p, line_spacing=28)
    run = p.add_run(str(copy_number))
    set_font(run, 'Times New Roman', 16)  # 三号


def _add_security_urgency_block(doc, header_data):
    """密级 / 紧急程度 (顶格版心左上角第二/三行) — 4 种格式共有"""
    security_level = header_data.get('security_level')
    if security_level:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, line_spacing=28)
        security_period = header_data.get('security_period', '')
        text = f'{security_level}★{security_period}年' if security_period else security_level
        run = p.add_run(text)
        set_font(run, '黑体', 16, bold=True)

    urgency = header_data.get('urgency')
    if urgency:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_spacing(p, line_spacing=28)
        run = p.add_run(urgency)
        set_font(run, '黑体', 16, bold=True)


def _add_issuer_mark(doc, text, font_pt=26, blank_lines=1):
    """发文机关标志 (红色小标宋体居中) — 通用实现"""
    if not text:
        return
    add_empty_lines(doc, blank_lines, 22)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=0, line_spacing=42)
    run = p.add_run(text)
    set_font(run, '小标宋体', font_pt, color=RGBColor(0xFF, 0x00, 0x00))


def _add_doc_number_centered(doc, doc_number):
    """发文字号居中 (下行文/平行文) — general 用"""
    doc_num_text = format_doc_number(
        doc_number.get('prefix', ''),
        doc_number.get('year', ''),
        doc_number.get('sequence', ''),
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=28)
    run = p.add_run(doc_num_text)
    set_font(run, '仿宋', 16)


def _add_doc_number_with_signers(doc, doc_number, signers):
    """发文字号居左+签发人居右 (上行文) — general 用"""
    doc_num_text = format_doc_number(
        doc_number.get('prefix', ''),
        doc_number.get('year', ''),
        doc_number.get('sequence', ''),
    )

    p = doc.add_paragraph()
    set_paragraph_spacing(p, line_spacing=28)
    run = p.add_run(f'\u3000{doc_num_text}')  # 发文字号左空一字
    set_font(run, '仿宋', 16)
    signer_text = '、'.join(signers)
    run2 = p.add_run(f'    签发人：{signer_text}')
    set_font(run2, '仿宋', 16)


def _add_order_number(doc, order_number):
    """令号居中 (命令格式 §4.2)"""
    if not order_number:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, line_spacing=28)
    run = p.add_run(order_number)
    set_font(run, '仿宋', 16)


def _add_attendees(doc, attendees):
    """出席人员名单 (纪要格式 §4.3)

    GB/T 9704 §4.3: 「'出席'二字后标全角冒号；冒号后用三号仿宋体字标注单位和姓名」
    整体作为主体的开头部分追加；放在主体 build_body 调用里更符合实际版面位置。
    """
    if not attendees:
        return
    # 主体内部处理：这里仅作为占位返回，由 build_minutes_body 调用
    return attendees


def build_header(doc, header_data, format_type='general'):
    """构建版头区域 — 根据 format_type 分派"""
    fmt = format_type if format_type in FORMAT_TYPES else 'general'

    # 份号 / 密级 / 紧急 (4 种格式共有)
    _add_copy_number(doc, header_data)
    _add_security_urgency_block(doc, header_data)

    issuer_mark = header_data.get('issuer_mark', '')
    doc_number = header_data.get('doc_number', {})
    signers = header_data.get('signer', [])

    if fmt == 'general':
        # 通用: 机关标志居中(26pt)，发文字号居中或居左+签发人
        _add_issuer_mark(doc, issuer_mark, font_pt=26)
        if doc_number:
            if signers:
                _add_doc_number_with_signers(doc, doc_number, signers)
            else:
                _add_doc_number_centered(doc, doc_number)
        # 红色分隔线
        add_empty_lines(doc, 1, 8)
        add_red_line(doc)

    elif fmt == 'letter':
        # 信函: §4.1 — 机关标志上距上页边 30mm，无发文字号，
        # 机关标志下 4mm 处印红色双线 (上粗下细)，距下页边 20mm 处印红色双线 (上细下粗)
        # 线长 170mm，居中排布
        _add_issuer_mark(doc, issuer_mark, font_pt=26, blank_lines=0)
        # 红色双线 (上粗下细)
        add_empty_lines(doc, 1, 4)
        add_red_double_line_top(doc)
        # 双线下空二行放主送机关/正文 — 由 build_letter_body 处理

    elif fmt == 'order':
        # 命令(令): §4.2 — 机关标志由全称加"命令"或"令"组成，上距版心 20mm
        # 发文机关标志下空二行居中编排令号
        order_mark = issuer_mark if issuer_mark else '×××命令'
        _add_issuer_mark(doc, order_mark, font_pt=26, blank_lines=0)
        # 令号 — 在 build_header_data 中读 order_number
        order_number = header_data.get('order_number', '')
        if order_number:
            add_empty_lines(doc, 2, 22)
            _add_order_number(doc, order_number)

    elif fmt == 'minutes':
        # 纪要: §4.3 — "×××××纪要" 标志 + 出席人员名单
        # 出席人员名单在主体 build_minutes_body 中追加
        minutes_mark = issuer_mark if issuer_mark else '×××××纪要'
        _add_issuer_mark(doc, minutes_mark, font_pt=26, blank_lines=0)

    return doc


def build_body(doc, subject_data, format_type='general'):
    """构建主体区域 — 根据 format_type 分派"""
    fmt = format_type if format_type in FORMAT_TYPES else 'general'
    if fmt == 'letter':
        return _build_letter_body(doc, subject_data)
    if fmt == 'order':
        return _build_order_body(doc, subject_data)
    if fmt == 'minutes':
        return _build_minutes_body(doc, subject_data)
    return _build_general_body(doc, subject_data)


def _build_general_body(doc, subject_data):
    """通用格式主体 — 标题居中, 主送机关, 正文 4 级层次, 附件, 署名, 附注"""

    # 标题 - 二号小标宋体，居中
    # 支持 title_lines 数组 (梯形/菱形排版)
    title_lines = subject_data.get('title_lines')
    title = subject_data.get('title', '')
    if title_lines:
        add_empty_lines(doc, 2, 22)  # 红色分隔线下空二行
        for line in title_lines:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, line_spacing=32)
            run = p.add_run(line)
            set_font(run, '小标宋体', 22)
    elif title:
        add_empty_lines(doc, 2, 22)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, line_spacing=32)
        run = p.add_run(title)
        set_font(run, '小标宋体', 22)

    # 主送机关 - 三号仿宋，顶格
    recipients = subject_data.get('recipients', [])
    if recipients:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=28)
        recipient_text = '、'.join(recipients) + '：'
        run = p.add_run(recipient_text)
        set_font(run, '仿宋', 16)

    _add_body_paragraphs(doc, subject_data.get('body', []))
    _add_attachment_block(doc, subject_data.get('attachments', []))
    _add_issuer_signature(doc, subject_data)
    return doc


def _build_letter_body(doc, subject_data):
    """信函格式主体 (§4.1)

    双线 (上粗下细) 在 build_header 中已加入；
    主体从这里开始：标题、主送机关、正文、署名、附注。
    注意：信函格式无版记 (由 build_footer 处理)。
    """
    return _build_general_body(doc, subject_data)


def _build_order_body(doc, subject_data):
    """命令(令)格式主体 (§4.2)

    令号在 build_header 中已加入；令号下空二行编排正文。
    命令(令) 可以有标题 (描述命令事项)；无主送机关。
    """
    add_empty_lines(doc, 2, 22)  # 令号下空二行
    # 标题 (如果有)：二号小标宋体，居中
    title = subject_data.get('title', '')
    if title:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p, line_spacing=32)
        run = p.add_run(title)
        set_font(run, '小标宋体', 22)
    _add_body_paragraphs(doc, subject_data.get('body', []))
    _add_issuer_signature(doc, subject_data, no_seal=True)
    return doc


def _build_minutes_body(doc, subject_data):
    """纪要格式主体 (§4.3)

    出席人员名单: 三号黑体字，正文或附件说明下空一行左空二字；
    "出席"二字后标全角冒号；冒号后用三号仿宋体字标注单位和姓名。
    """
    _build_general_body(doc, subject_data)
    # 出席人员名单: 在主体 build_general_body 之后追加 (避免破坏签名/版记顺序)
    attendees = subject_data.get('attendees', [])
    if attendees:
        # 出席标记
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = Pt(32)
        set_paragraph_spacing(p, line_spacing=28)
        run = p.add_run('出席：')
        set_font(run, '黑体', 16, bold=True)
        # 名单 (三号仿宋) — 通常是一个字符串或数组
        if isinstance(attendees, str):
            attendees_text = attendees
        else:
            attendees_text = '、'.join(attendees)
        run2 = p.add_run(attendees_text)
        set_font(run2, '仿宋', 16)
    return doc


def _add_body_paragraphs(doc, body):
    """正文段落 (4 级层次结构)"""
    for item in body:
        if isinstance(item, str):
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.first_line_indent = Pt(32)
            set_paragraph_spacing(p, line_spacing=28)
            run = p.add_run(item)
            set_font(run, '仿宋', 16)
        elif isinstance(item, dict):
            level = item.get('level', 1)
            text = item.get('text', '')
            paragraphs = item.get('paragraphs', [])

            font_name = LEVEL_FONT_MAP.get(level, '仿宋')
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.first_line_indent = Pt(32)
            set_paragraph_spacing(p, line_spacing=28)
            run = p.add_run(text)
            set_font(run, font_name, 16)

            for para_text in paragraphs:
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.first_line_indent = Pt(32)
                set_paragraph_spacing(p, line_spacing=28)
                run = p.add_run(para_text)
                set_font(run, '仿宋', 16)


def _add_attachment_block(doc, attachments):
    """附件说明 — 多个附件用 <w:br/> 分行 (GB/T 9704 §3.2.4)"""
    if not attachments:
        return
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.first_line_indent = Pt(32)
    set_paragraph_spacing(p, line_spacing=28)
    run = p.add_run('附件：')
    set_font(run, '仿宋', 16)
    for i, att in enumerate(attachments, 1):
        att_name = att.get('name', f'附件{i}') if isinstance(att, dict) else str(att)
        if i > 1:
            run.add_break()  # <w:br/> — 显式换行，不依赖 python-docx 的隐式 '\n' 处理
        att_run = p.add_run(f'{i}.{att_name}')
        set_font(att_run, '仿宋', 16)


def _add_issuer_signature(doc, subject_data, no_seal=False):
    """发文机关署名 + 成文日期 + 附注

    支持单署名 (issuer_name) 与联合行文 (issuer_names 数组)。
    """
    issuer_name = subject_data.get('issuer_name', '')
    issuer_names = subject_data.get('issuer_names')  # 联合行文: 主办机关在前
    issue_date = subject_data.get('issue_date', '')
    seal = subject_data.get('seal', True) and not no_seal

    if not (issuer_name or issuer_names or issue_date):
        return

    add_empty_lines(doc, 1, 16)

    if issuer_names:
        # 联合行文: 主办机关在前，每个机关一行
        for name in issuer_names:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_paragraph_spacing(p, line_spacing=28)
            run = p.add_run(name)
            set_font(run, '仿宋', 16)
    elif issuer_name:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(p, line_spacing=28)
        run = p.add_run(issuer_name)
        set_font(run, '仿宋', 16)

    if issue_date:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_paragraph_spacing(p, line_spacing=28)
        run = p.add_run(f'{issue_date}    ')
        set_font(run, '仿宋', 16)

    notes = subject_data.get('notes')
    if notes:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = Pt(32)
        set_paragraph_spacing(p, line_spacing=28)
        run = p.add_run(f'（{notes}）')
        set_font(run, '仿宋', 16)


def build_footer(doc, footer_data, format_type='general'):
    """构建版记区域 — 信函 (§4.1) 不加版记的印发机关/印发日期与分隔线"""
    fmt = format_type if format_type in FORMAT_TYPES else 'general'

    if fmt == 'letter':
        # 信函: 无版记，但页边 20mm 处印红色双线 (上细下粗) 作为版尾装饰
        # 在主体结束后添加，build_body 之后再调到这里
        add_red_double_line_bottom(doc)
        return doc

    if not footer_data:
        return doc

    cc_list = footer_data.get('cc', [])
    print_org = footer_data.get('print_org', '')
    print_date = footer_data.get('print_date', '')

    if not cc_list and not print_org:
        return doc

    # 版记首条粗分隔线
    add_thick_line(doc)

    # 抄送机关
    if cc_list:
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.first_line_indent = Pt(28)
        set_paragraph_spacing(p, line_spacing=24)
        cc_text = '抄送：' + '、'.join(cc_list) + '。'
        run = p.add_run(cc_text)
        set_font(run, '仿宋', 14)

    if cc_list and print_org:
        add_thin_line(doc)

    if print_org or print_date:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, line_spacing=24)
        if print_org:
            run = p.add_run(f' {print_org}')
            set_font(run, '仿宋', 14)
        if print_date:
            remaining = 28 - len(print_org or '') - len(print_date) - 2
            spaces = ' ' * max(remaining, 1)
            run = p.add_run(f'{spaces}{print_date} 印发 ')
            set_font(run, '仿宋', 14)

    add_thick_line(doc)
    return doc


def _enable_even_odd_headers(doc):
    """在 settings.xml 中启用奇偶页不同 header/footer (GB/T 9704 §3.4)"""
    settings = doc.settings.element
    # 移除已有
    existing = settings.find(qn('w:evenAndOddHeaders'))
    if existing is not None:
        settings.remove(existing)
    el = parse_xml(f'<w:evenAndOddHeaders {nsdecls("w")}/>')
    settings.insert(0, el)


def _add_page_field(p, font_name='宋体', font_size_pt=14):
    """在段落中添加 PAGE 域 — 四号半角宋体阿拉伯数字"""
    run = p.add_run()
    set_font(run, font_name, font_size_pt)
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)


def _add_dash_lines(p, side='right'):
    """在 PAGE 数字旁添加一字线 (—)

    side='right' (奇数页):   —X—   居右
    side='left'  (偶数页):   —X—   居左
    """
    pass  # 简化：实际渲染时把 — 直接加到 page 段落文本里


def add_page_numbers(doc, format_type='general'):
    """添加页码 — GB/T 9704 §3.4:
    四号半角宋体阿拉伯数字，单页码居右空一字，双页码居左空一字；
    一字线上距版心下边缘 7mm；空白页和版记页均不编排页码。

    - general/minutes/order: 启用奇偶页不同 footer
    - letter (§4.1): 首页不显示页码 (通过第一页 footer 留空实现)
    """
    fmt = format_type if format_type in FORMAT_TYPES else 'general'

    if fmt == 'letter':
        # 信函 (§4.1): 首页不显示页码，但其他页仍编页码 — 启用 first-page footer
        # first_page_footer 留空；odd/even footer 正常加页码
        for section in doc.sections:
            section.different_first_page_header_footer = True
            fpf = section.first_page_footer
            fpf.is_linked_to_previous = False
            # 清空 first-page footer (留空)
            for p in list(fpf.paragraphs):
                for r in list(p.runs):
                    r._element.getparent().remove(r._element)
        # 继续走 odd/even footer 路径
        _enable_even_odd_headers(doc)
        for section in doc.sections:
            odd_footer = section.footer
            odd_footer.is_linked_to_previous = False
            op = odd_footer.paragraphs[0] if odd_footer.paragraphs else odd_footer.add_paragraph()
            op.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in list(op.runs):
                r._element.getparent().remove(r._element)
            op.add_run('—').font.size = Pt(14)
            _add_page_field(op, '宋体', 14)
            op.add_run('—').font.size = Pt(14)

            even_footer = section.even_page_footer
            even_footer.is_linked_to_previous = False
            ep = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
            ep.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in list(ep.runs):
                r._element.getparent().remove(r._element)
            ep.add_run('—').font.size = Pt(14)
            _add_page_field(ep, '宋体', 14)
            ep.add_run('—').font.size = Pt(14)
        return

    # 启用奇偶页不同 footer
    _enable_even_odd_headers(doc)

    for section in doc.sections:
        # 奇数页 footer: 居右 (—X—)
        odd_footer = section.footer
        odd_footer.is_linked_to_previous = False
        op = odd_footer.paragraphs[0] if odd_footer.paragraphs else odd_footer.add_paragraph()
        op.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        # 清空再写入
        for r in list(op.runs):
            r._element.getparent().remove(r._element)
        # —X—
        run_l = op.add_run('—')
        set_font(run_l, '宋体', 14)
        _add_page_field(op, '宋体', 14)
        run_r = op.add_run('—')
        set_font(run_r, '宋体', 14)

        # 偶数页 footer: 居左
        even_footer = section.even_page_footer
        even_footer.is_linked_to_previous = False
        ep = even_footer.paragraphs[0] if even_footer.paragraphs else even_footer.add_paragraph()
        ep.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for r in list(ep.runs):
            r._element.getparent().remove(r._element)
        run_l = ep.add_run('—')
        set_font(run_l, '宋体', 14)
        _add_page_field(ep, '宋体', 14)
        run_r = ep.add_run('—')
        set_font(run_r, '宋体', 14)


def generate_gov_document(input_data, output_path):
    """主函数：生成公文"""
    doc = Document()

    header_data = input_data.get('header', {})
    subject_data = input_data.get('subject', {})
    footer_data = input_data.get('footer', {})
    format_type = input_data.get('format_type', input_data.get('doc_type', 'general'))
    show_page_number = input_data.get('page_number', True)

    if format_type not in FORMAT_TYPES:
        print(f'警告: format_type={format_type!r} 不在 {FORMAT_TYPES} 内，按 general 处理', file=sys.stderr)
        format_type = 'general'

    # 1. 页面设置
    setup_page(doc, format_type)

    # 2. 构建版头
    build_header(doc, header_data, format_type)

    # 3. 构建主体
    build_body(doc, subject_data, format_type)

    # 4. 构建版记 (含信函格式的下边红色双线)
    build_footer(doc, footer_data, format_type)

    # 5. 添加页码
    if show_page_number:
        add_page_numbers(doc, format_type)

    # 6. 保存
    doc.save(output_path)
    print(f'公文已生成: {output_path}')
    return output_path


def main():
    parser = argparse.ArgumentParser(description='党政机关公文生成器 (GB/T 9704-2012)')
    parser.add_argument('--input', '-i', required=True, help='JSON输入文件路径')
    parser.add_argument('--output', '-o', required=True, help='输出docx文件路径')
    args = parser.parse_args()

    with open(args.input, 'r', encoding='utf-8') as f:
        input_data = json.load(f)

    output_path = args.output
    if not output_path.endswith('.docx'):
        output_path += '.docx'

    generate_gov_document(input_data, output_path)


if __name__ == '__main__':
    main()